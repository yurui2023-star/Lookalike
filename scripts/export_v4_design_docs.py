#!/usr/bin/env python3
"""Export Lookalike v4.0 detailed design HTML to Markdown and Word."""

from __future__ import annotations

import re
import shutil
from pathlib import Path

from bs4 import BeautifulSoup, NavigableString, Tag
from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor

ROOT = Path(__file__).resolve().parents[1]
HTML_PATH = ROOT / "docs" / "Lookalike_Detailed_Design_v4.0_BatchScoring.html"
MD_NAME = "Lookalike_Detailed_Design_v4.0_BatchScoring.md"
DOCX_NAME = "Lookalike_Detailed_Design_v4.0_BatchScoring.docx"

GREEN = RGBColor(0x0F, 0x6E, 0x56)
DARK = RGBColor(0x1A, 0x1A, 0x2E)
MUTED = RGBColor(0x6B, 0x72, 0x80)
AMBER = RGBColor(0x85, 0x4F, 0x0B)
RED = RGBColor(0xA3, 0x2D, 0x2D)
BLUE = RGBColor(0x18, 0x5F, 0xA5)
WHITE = RGBColor(0xFF, 0xFF, 0xFF)
HEADER_FILL = "0F6E56"
TH_FILL = "F0F4F2"
CALLOUT_FILLS = {
    "ok": ("E1F5EE", GREEN),
    "warn": ("FAEEDA", AMBER),
    "info": ("E6F1FB", BLUE),
    "danger": ("FCEBEB", RED),
}
CALLOUT_LABEL = {"ok": "确认", "warn": "注意", "info": "说明", "danger": "禁止"}


def inline_md(el: Tag | NavigableString | None) -> str:
    if el is None:
        return ""
    if isinstance(el, NavigableString):
        return str(el).replace("\n", " ")
    parts: list[str] = []
    for child in el.children:
        if isinstance(child, NavigableString):
            parts.append(str(child).replace("\n", " "))
        elif isinstance(child, Tag):
            name = child.name
            inner = inline_md(child).strip()
            if name in {"strong", "b"}:
                parts.append(f"**{inner}**" if inner else "")
            elif name == "code":
                parts.append(f"`{child.get_text()}`")
            elif name == "br":
                parts.append("\n")
            elif name == "span" and "badge" in child.get("class", []):
                parts.append(inner)
            else:
                parts.append(inline_md(child))
    return re.sub(r"[ \t]{2,}", " ", "".join(parts))


def inline_runs(el: Tag | NavigableString | None) -> list[tuple[str, dict]]:
    if el is None:
        return []
    if isinstance(el, NavigableString):
        text = str(el).replace("\n", " ")
        return [(text, {})] if text else []
    runs: list[tuple[str, dict]] = []
    for child in el.children:
        if isinstance(child, NavigableString):
            text = str(child).replace("\n", " ")
            if text:
                runs.append((text, {}))
        elif isinstance(child, Tag):
            name = child.name
            if name == "br":
                runs.append(("\n", {}))
            elif name == "code":
                runs.append((child.get_text(), {"code": True}))
            elif name in {"strong", "b"}:
                for text, flags in inline_runs(child):
                    flags = dict(flags)
                    flags["strong"] = True
                    runs.append((text, flags))
            else:
                runs.extend(inline_runs(child))
    return runs


def callout_kind(classes: list[str]) -> str:
    for k in ("ok", "warn", "danger", "info"):
        if f"callout-{k}" in classes:
            return k
    return "info"


def parse(path: Path) -> dict:
    soup = BeautifulSoup(path.read_text(encoding="utf-8"), "lxml")
    header = soup.select_one(".doc-header")
    title = header.find("h1").get_text(strip=True) if header else ""
    subtitle = header.select_one(".subtitle").get_text(strip=True) if header else ""
    meta: list[tuple[str, str]] = []
    if header:
        for item in header.select(".doc-meta > div"):
            label = item.select_one(".label")
            value = item.find("strong")
            if label and value:
                meta.append((label.get_text(strip=True), value.get_text(strip=True)))

    blocks: list[dict] = []
    container = soup.select_one(".container")
    assert container is not None
    for node in container.children:
        if not isinstance(node, Tag):
            continue
        classes = node.get("class", [])
        if "doc-header" in classes:
            continue
        if "toc" in classes:
            items = [li.get_text(strip=True) for li in node.select("ol > li")]
            blocks.append({"type": "toc", "items": items})
            continue
        if node.name == "section":
            for child in node.children:
                if not isinstance(child, Tag):
                    continue
                blocks.extend(_convert(child))
            continue
        if node.name == "p" and "footer" in classes:
            blocks.append({"type": "p", "md": inline_md(node).strip(), "runs": inline_runs(node)})
    return {"title": title, "subtitle": subtitle, "meta": meta, "blocks": blocks}


def _convert(node: Tag) -> list[dict]:
    name = node.name
    classes = node.get("class", [])
    if name in {"h2", "h3", "h4"}:
        return [{"type": "h", "level": int(name[1]), "text": node.get_text(strip=True)}]
    if "callout" in classes:
        kind = callout_kind(classes)
        return [{"type": "callout", "kind": kind, "md": inline_md(node).strip(), "runs": inline_runs(node)}]
    if "mermaid" in classes:
        return [{"type": "mermaid", "text": node.get_text().strip()}]
    if name == "pre":
        return [{"type": "code", "text": node.get_text().strip("\n")}]
    if name == "table":
        rows = []
        for tr in node.select("tr"):
            cells = [inline_runs(cell) for cell in tr.find_all(["th", "td"], recursive=False)]
            rows.append(cells)
        return [{"type": "table", "rows": rows}]
    if name in {"ul", "ol"}:
        items = [{"md": inline_md(li).strip(), "runs": inline_runs(li)} for li in node.find_all("li", recursive=False)]
        return [{"type": "list", "kind": name, "items": items}]
    if name == "p":
        return [{"type": "p", "md": inline_md(node).strip(), "runs": inline_runs(node), "strike": "strike" in classes}]
    if "stat-row" in classes:
        items = []
        for stat in node.select(".stat"):
            num = stat.select_one(".num")
            label = stat.select_one(".label")
            items.append((num.get_text(strip=True) if num else "", label.get_text(strip=True) if label else ""))
        return [{"type": "stats", "items": items}]
    if name == "div":
        out: list[dict] = []
        for child in node.children:
            if isinstance(child, Tag):
                out.extend(_convert(child))
        return out
    return []


def runs_plain(runs: list[tuple[str, dict]]) -> str:
    return "".join(t for t, _ in runs).strip()


def to_markdown(doc: dict) -> str:
    lines: list[str] = [f"# {doc['title']}", "", f"> {doc['subtitle']}", ""]
    if doc["meta"]:
        lines += ["| 项 | 内容 |", "| --- | --- |"]
        for k, v in doc["meta"]:
            lines.append(f"| {k} | {v} |")
        lines.append("")
    for block in doc["blocks"]:
        t = block["type"]
        if t == "toc":
            lines += ["## 目录", ""]
            for item in block["items"]:
                lines.append(f"- {item}")
            lines.append("")
        elif t == "h":
            lines += ["", f"{'#' * block['level']} {block['text']}", ""]
        elif t == "p":
            text = block["md"].strip()
            if block.get("strike"):
                text = f"~~{text}~~"
            if text:
                lines += [text, ""]
        elif t == "callout":
            label = CALLOUT_LABEL[block["kind"]]
            lines += [f"> **{label}：** {block['md']}", ""]
        elif t == "list":
            for i, item in enumerate(block["items"], 1):
                bullet = f"{i}." if block["kind"] == "ol" else "-"
                lines.append(f"{bullet} {item['md']}")
            lines.append("")
        elif t == "table":
            rows = [[runs_plain(c) for c in row] for row in block["rows"]]
            if not rows:
                continue
            width = max(len(r) for r in rows)
            rows = [r + [""] * (width - len(r)) for r in rows]
            lines.append("| " + " | ".join(_escape_md_cell(c) for c in rows[0]) + " |")
            lines.append("| " + " | ".join("---" for _ in rows[0]) + " |")
            for row in rows[1:]:
                lines.append("| " + " | ".join(_escape_md_cell(c) for c in row) + " |")
            lines.append("")
        elif t == "code":
            lines += ["```", block["text"], "```", ""]
        elif t == "mermaid":
            lines += ["```mermaid", block["text"], "```", ""]
        elif t == "stats":
            for num, label in block["items"]:
                lines.append(f"- **{num}**：{label}")
            lines.append("")
    text = re.sub(r"\n{3,}", "\n\n", "\n".join(lines)).strip() + "\n"
    return text


def _escape_md_cell(text: str) -> str:
    return text.replace("|", "\\|").replace("\n", " ")


def set_run_font(run, *, name="Calibri", east="微软雅黑", size=None, bold=None, color=None):
    run.font.name = name
    rPr = run._element.get_or_add_rPr()
    rFonts = rPr.find(qn("w:rFonts"))
    if rFonts is None:
        rFonts = OxmlElement("w:rFonts")
        rPr.append(rFonts)
    rFonts.set(qn("w:ascii"), name)
    rFonts.set(qn("w:hAnsi"), name)
    rFonts.set(qn("w:eastAsia"), east)
    if size is not None:
        run.font.size = size
    if bold is not None:
        run.bold = bold
    if color is not None:
        run.font.color.rgb = color


def shade_cell(cell, fill: str) -> None:
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    shd.set(qn("w:val"), "clear")
    tcPr.append(shd)


def set_cell_border(cell) -> None:
    tcPr = cell._tc.get_or_add_tcPr()
    tcBorders = OxmlElement("w:tcBorders")
    for edge in ("top", "left", "bottom", "right"):
        el = OxmlElement(f"w:{edge}")
        el.set(qn("w:val"), "single")
        el.set(qn("w:sz"), "4")
        el.set(qn("w:color"), "D1D5DB")
        tcBorders.append(el)
    tcPr.append(tcBorders)


def add_runs(paragraph, runs, *, size=Pt(11), color=DARK) -> None:
    for text, flags in runs:
        if not text:
            continue
        run = paragraph.add_run(text)
        is_code = bool(flags.get("code"))
        set_run_font(
            run,
            name="Consolas" if is_code else "Calibri",
            size=size,
            bold=bool(flags.get("strong")),
            color=BLUE if is_code else color,
        )
        if is_code:
            rPr = run._element.get_or_add_rPr()
            rFonts = rPr.find(qn("w:rFonts"))
            if rFonts is not None:
                rFonts.set(qn("w:ascii"), "Consolas")
                rFonts.set(qn("w:hAnsi"), "Consolas")
        if flags.get("strike"):
            run.font.strike = True


def add_horizontal_line(paragraph) -> None:
    pPr = paragraph._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "12")
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), "0F6E56")
    pBdr.append(bottom)
    pPr.append(pBdr)


def add_page_field(paragraph) -> None:
    fld1 = OxmlElement("w:fldChar")
    fld1.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = " PAGE "
    fld2 = OxmlElement("w:fldChar")
    fld2.set(qn("w:fldCharType"), "end")
    r = paragraph.add_run()
    r._r.append(fld1)
    r2 = paragraph.add_run()
    r2._r.append(instr)
    r3 = paragraph.add_run()
    r3._r.append(fld2)
    for rr in (r, r2, r3):
        set_run_font(rr, size=Pt(8), color=MUTED)


def to_docx(doc: dict, path: Path) -> None:
    document = Document()
    section = document.sections[0]
    section.page_width = Cm(21.0)
    section.page_height = Cm(29.7)
    section.left_margin = Cm(2.2)
    section.right_margin = Cm(2.2)
    section.top_margin = Cm(2.0)
    section.bottom_margin = Cm(2.0)

    hp = section.header.paragraphs[0]
    hp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = hp.add_run("MB Bank Lookalike  ·  详细设计说明书 v4.0")
    set_run_font(run, size=Pt(8), color=MUTED)

    fp = section.footer.paragraphs[0]
    fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = fp.add_run("对齐 SAD V1.0 与 Sales CDP 附件 schema  ·  2026-08-17  ·  第 ")
    set_run_font(run, size=Pt(8), color=MUTED)
    add_page_field(fp)
    run = fp.add_run(" 页")
    set_run_font(run, size=Pt(8), color=MUTED)

    p = document.add_paragraph()
    p.paragraph_format.space_after = Pt(6)
    run = p.add_run(doc["title"])
    set_run_font(run, size=Pt(20), bold=True, color=GREEN)

    p = document.add_paragraph()
    p.paragraph_format.space_after = Pt(12)
    run = p.add_run(doc["subtitle"])
    set_run_font(run, size=Pt(11), color=MUTED)
    add_horizontal_line(p)

    if doc["meta"]:
        table = document.add_table(rows=len(doc["meta"]), cols=2)
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        for i, (k, v) in enumerate(doc["meta"]):
            c0, c1 = table.rows[i].cells
            c0.text = ""
            c1.text = ""
            r0 = c0.paragraphs[0].add_run(k)
            set_run_font(r0, size=Pt(10), bold=True, color=WHITE)
            r1 = c1.paragraphs[0].add_run(v)
            set_run_font(r1, size=Pt(10), color=DARK)
            shade_cell(c0, HEADER_FILL)
            shade_cell(c1, "F7F8FA")
            set_cell_border(c0)
            set_cell_border(c1)
        document.add_paragraph()

    for block in doc["blocks"]:
        t = block["type"]
        if t == "toc":
            p = document.add_paragraph()
            p.paragraph_format.space_before = Pt(12)
            run = p.add_run("目录")
            set_run_font(run, size=Pt(16), bold=True, color=GREEN)
            add_horizontal_line(p)
            for item in block["items"]:
                bp = document.add_paragraph(style="List Number")
                run = bp.add_run(re.sub(r"^\d+\.\s*", "", item))
                set_run_font(run, size=Pt(11), color=DARK)
        elif t == "h":
            p = document.add_paragraph()
            p.paragraph_format.space_before = Pt(16 if block["level"] == 2 else 10)
            p.paragraph_format.space_after = Pt(6)
            size = {2: 16, 3: 13, 4: 12}.get(block["level"], 12)
            run = p.add_run(block["text"])
            set_run_font(run, size=Pt(size), bold=True, color=GREEN if block["level"] == 2 else DARK)
            if block["level"] == 2:
                add_horizontal_line(p)
        elif t == "p":
            p = document.add_paragraph()
            p.paragraph_format.space_after = Pt(8)
            p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
            add_runs(p, block["runs"])
            if block.get("strike"):
                for run in p.runs:
                    run.font.strike = True
        elif t == "callout":
            table = document.add_table(rows=1, cols=1)
            cell = table.cell(0, 0)
            fill, color = CALLOUT_FILLS[block["kind"]]
            shade_cell(cell, fill)
            set_cell_border(cell)
            para = cell.paragraphs[0]
            run = para.add_run(f"{CALLOUT_LABEL[block['kind']]}  ")
            set_run_font(run, size=Pt(10), bold=True, color=color)
            add_runs(para, block["runs"], size=Pt(10), color=DARK)
            document.add_paragraph()
        elif t == "list":
            style = "List Number" if block["kind"] == "ol" else "List Bullet"
            for item in block["items"]:
                p = document.add_paragraph(style=style)
                add_runs(p, item["runs"], size=Pt(11))
        elif t == "table":
            rows = block["rows"]
            if not rows:
                continue
            width = max(len(r) for r in rows)
            table = document.add_table(rows=len(rows), cols=width)
            table.alignment = WD_TABLE_ALIGNMENT.CENTER
            for i, row in enumerate(rows):
                for j in range(width):
                    cell = table.cell(i, j)
                    cell.text = ""
                    runs = row[j] if j < len(row) else []
                    para = cell.paragraphs[0]
                    add_runs(para, runs, size=Pt(9), color=DARK)
                    set_cell_border(cell)
                    if i == 0:
                        shade_cell(cell, TH_FILL)
                        for run in para.runs:
                            run.bold = True
                            run.font.color.rgb = GREEN
            document.add_paragraph()
        elif t == "code":
            p = document.add_paragraph()
            p.paragraph_format.space_before = Pt(4)
            p.paragraph_format.space_after = Pt(10)
            run = p.add_run(block["text"])
            set_run_font(run, name="Consolas", size=Pt(8), color=RGBColor(0x1E, 0x24, 0x30))
        elif t == "mermaid":
            p = document.add_paragraph()
            run = p.add_run("【架构图 / 时序图 — Mermaid 源】")
            set_run_font(run, size=Pt(9), bold=True, color=MUTED)
            p = document.add_paragraph()
            run = p.add_run(block["text"])
            set_run_font(run, name="Consolas", size=Pt(8), color=DARK)
        elif t == "stats":
            n = max(1, len(block["items"]))
            table = document.add_table(rows=1, cols=n)
            for i, (num, label) in enumerate(block["items"]):
                cell = table.cell(0, i)
                cell.text = ""
                shade_cell(cell, "E1F5EE")
                set_cell_border(cell)
                p = cell.paragraphs[0]
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                run = p.add_run(num + "\n")
                set_run_font(run, size=Pt(14), bold=True, color=GREEN)
                run = p.add_run(label)
                set_run_font(run, size=Pt(9), color=MUTED)
            document.add_paragraph()

    document.save(path)


def desktop_dir() -> Path:
    desktop = Path.home() / "Desktop"
    desktop.mkdir(parents=True, exist_ok=True)
    return desktop


def main() -> None:
    parsed = parse(HTML_PATH)
    md_text = to_markdown(parsed)

    md_repo = ROOT / "docs" / MD_NAME
    md_repo.write_text(md_text, encoding="utf-8")

    desktop = desktop_dir()
    md_desktop = desktop / MD_NAME
    docx_desktop = desktop / DOCX_NAME
    md_desktop.write_text(md_text, encoding="utf-8")
    to_docx(parsed, docx_desktop)

    artifacts = Path("/opt/cursor/artifacts")
    artifacts.mkdir(parents=True, exist_ok=True)
    shutil.copy2(md_desktop, artifacts / MD_NAME)
    shutil.copy2(docx_desktop, artifacts / DOCX_NAME)
    shutil.copy2(docx_desktop, ROOT / "docs" / DOCX_NAME)

    h2 = [b["text"] for b in parsed["blocks"] if b["type"] == "h" and b["level"] == 2]
    print(f"title: {parsed['title']}")
    print(f"meta: {parsed['meta']}")
    print(f"h2 ({len(h2)}): {h2}")
    print(f"blocks: {len(parsed['blocks'])}")
    print(f"markdown: {md_desktop} ({md_desktop.stat().st_size} bytes)")
    print(f"word:     {docx_desktop} ({docx_desktop.stat().st_size} bytes)")


if __name__ == "__main__":
    main()
