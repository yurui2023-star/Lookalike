#!/usr/bin/env python3
"""Convert the clean v4.0 Markdown detailed design to Word and copy to Desktop."""

from __future__ import annotations

import re
import shutil
from pathlib import Path

from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor

ROOT = Path(__file__).resolve().parents[1]
MD_PATH = ROOT / "docs" / "Lookalike_Detailed_Design_v4.0_BatchScoring.md"
DOCX_NAME = "Lookalike_Detailed_Design_v4.0_BatchScoring.docx"

GREEN = RGBColor(0x0F, 0x6E, 0x56)
DARK = RGBColor(0x1A, 0x1A, 0x2E)
MUTED = RGBColor(0x6B, 0x72, 0x80)
BLUE = RGBColor(0x18, 0x5F, 0xA5)
WHITE = RGBColor(0xFF, 0xFF, 0xFF)
HEADER_FILL = "0F6E56"
TH_FILL = "F0F4F2"


def set_run_font(run, *, name="Calibri", east="微软雅黑", size=None, bold=None, color=None):
    run.font.name = name
    rPr = run._element.get_or_add_rPr()
    rFonts = rPr.find(qn("w:rFonts"))
    if rFonts is None:
        rFonts = OxmlElement("w:rFonts")
        rPr.append(rFonts)
    rFonts.set(qn("w:ascii"), name)
    rFonts.set(qn("w:hAnsi"), name)
    rFonts.set(qn("w:eastAsia"), east)
    if size is not None:
        run.font.size = size
    if bold is not None:
        run.bold = bold
    if color is not None:
        run.font.color.rgb = color


def shade_cell(cell, fill: str) -> None:
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    shd.set(qn("w:val"), "clear")
    tcPr.append(shd)


def set_cell_border(cell) -> None:
    tcPr = cell._tc.get_or_add_tcPr()
    tcBorders = OxmlElement("w:tcBorders")
    for edge in ("top", "left", "bottom", "right"):
        el = OxmlElement(f"w:{edge}")
        el.set(qn("w:val"), "single")
        el.set(qn("w:sz"), "4")
        el.set(qn("w:color"), "D1D5DB")
        tcBorders.append(el)
    tcPr.append(tcBorders)


def add_horizontal_line(paragraph) -> None:
    pPr = paragraph._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "12")
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), "0F6E56")
    pBdr.append(bottom)
    pPr.append(pBdr)


def add_page_field(paragraph) -> None:
    fld1 = OxmlElement("w:fldChar")
    fld1.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = " PAGE "
    fld2 = OxmlElement("w:fldChar")
    fld2.set(qn("w:fldCharType"), "end")
    r = paragraph.add_run()
    r._r.append(fld1)
    r2 = paragraph.add_run()
    r2._r.append(instr)
    r3 = paragraph.add_run()
    r3._r.append(fld2)
    for rr in (r, r2, r3):
        set_run_font(rr, size=Pt(8), color=MUTED)


TOKEN = re.compile(r"(\*\*[^*]+\*\*|`[^`]+`)")


def add_inline(paragraph, text: str, *, size=Pt(11), color=DARK) -> None:
    for part in TOKEN.split(text):
        if not part:
            continue
        if part.startswith("**") and part.endswith("**"):
            run = paragraph.add_run(part[2:-2])
            set_run_font(run, size=size, bold=True, color=color)
        elif part.startswith("`") and part.endswith("`"):
            run = paragraph.add_run(part[1:-1])
            set_run_font(run, name="Consolas", size=size, color=BLUE)
        else:
            run = paragraph.add_run(part)
            set_run_font(run, size=size, color=color)


def split_row(line: str) -> list[str]:
    line = line.strip().strip("|")
    return [c.strip().replace("\\|", "|") for c in re.split(r"(?<!\\)\|", line)]


def is_sep(line: str) -> bool:
    cells = split_row(line)
    return bool(cells) and all(re.fullmatch(r":?-{3,}:?", c) for c in cells)


def md_to_docx(md_text: str, path: Path) -> None:
    document = Document()
    section = document.sections[0]
    section.page_width = Cm(21.0)
    section.page_height = Cm(29.7)
    section.left_margin = Cm(2.2)
    section.right_margin = Cm(2.2)
    section.top_margin = Cm(2.0)
    section.bottom_margin = Cm(2.0)

    hp = section.header.paragraphs[0]
    hp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = hp.add_run("MB Bank Lookalike  ·  详细设计说明书 v4.0")
    set_run_font(run, size=Pt(8), color=MUTED)

    fp = section.footer.paragraphs[0]
    fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = fp.add_run("Lookalike 月度批处理打分服务  ·  2026-08-17  ·  第 ")
    set_run_font(run, size=Pt(8), color=MUTED)
    add_page_field(fp)
    run = fp.add_run(" 页")
    set_run_font(run, size=Pt(8), color=MUTED)

    lines = md_text.splitlines()
    i = 0
    first_h1 = True
    while i < len(lines):
        line = lines[i]
        if not line.strip():
            i += 1
            continue
        if line.startswith("```"):
            lang = line[3:].strip()
            buf: list[str] = []
            i += 1
            while i < len(lines) and not lines[i].startswith("```"):
                buf.append(lines[i])
                i += 1
            i += 1
            p = document.add_paragraph()
            p.paragraph_format.space_after = Pt(10)
            label = "JSON" if lang == "json" else "SQL" if lang == "sql" else None
            if label:
                run = p.add_run(label + "\n")
                set_run_font(run, size=Pt(8), bold=True, color=MUTED)
            run = p.add_run("\n".join(buf))
            set_run_font(run, name="Consolas", size=Pt(8), color=RGBColor(0x1E, 0x24, 0x30))
            continue
        if line.startswith("|"):
            rows = [split_row(line)]
            i += 1
            if i < len(lines) and is_sep(lines[i]):
                i += 1
            while i < len(lines) and lines[i].startswith("|"):
                rows.append(split_row(lines[i]))
                i += 1
            width = max(len(r) for r in rows)
            table = document.add_table(rows=len(rows), cols=width)
            table.alignment = WD_TABLE_ALIGNMENT.CENTER
            for r_i, row in enumerate(rows):
                for c_i in range(width):
                    cell = table.cell(r_i, c_i)
                    cell.text = ""
                    para = cell.paragraphs[0]
                    add_inline(para, row[c_i] if c_i < len(row) else "", size=Pt(9))
                    set_cell_border(cell)
                    if r_i == 0:
                        shade_cell(cell, TH_FILL)
                        for run in para.runs:
                            run.bold = True
                            run.font.color.rgb = GREEN
            document.add_paragraph()
            continue
        m = re.match(r"^(#{1,4})\s+(.*)$", line)
        if m:
            level = len(m.group(1))
            text = m.group(2).strip()
            p = document.add_paragraph()
            if level == 1 and first_h1:
                first_h1 = False
                p.paragraph_format.space_after = Pt(10)
                run = p.add_run(text)
                set_run_font(run, size=Pt(20), bold=True, color=GREEN)
                add_horizontal_line(p)
            else:
                p.paragraph_format.space_before = Pt(16 if level == 2 else 10)
                p.paragraph_format.space_after = Pt(6)
                size = {2: 16, 3: 13, 4: 12}.get(level, 12)
                run = p.add_run(text)
                set_run_font(run, size=Pt(size), bold=True, color=GREEN if level == 2 else DARK)
                if level == 2:
                    add_horizontal_line(p)
            i += 1
            continue
        ol = re.match(r"^(\d+)\.\s+(.*)$", line)
        if ol:
            p = document.add_paragraph(style="List Number")
            add_inline(p, ol.group(2))
            i += 1
            continue
        if line.startswith("- "):
            p = document.add_paragraph(style="List Bullet")
            add_inline(p, line[2:])
            i += 1
            continue
        p = document.add_paragraph()
        p.paragraph_format.space_after = Pt(8)
        add_inline(p, line.strip())
        i += 1

    document.save(path)


def main() -> None:
    md_text = MD_PATH.read_text(encoding="utf-8")
    desktop = Path.home() / "Desktop"
    desktop.mkdir(parents=True, exist_ok=True)
    artifacts = Path("/opt/cursor/artifacts")
    artifacts.mkdir(parents=True, exist_ok=True)

    docx_repo = ROOT / "docs" / DOCX_NAME
    md_to_docx(md_text, docx_repo)

    md_desktop = desktop / MD_PATH.name
    docx_desktop = desktop / DOCX_NAME
    md_desktop.write_text(md_text, encoding="utf-8")
    shutil.copy2(docx_repo, docx_desktop)
    shutil.copy2(md_desktop, artifacts / MD_PATH.name)
    shutil.copy2(docx_desktop, artifacts / DOCX_NAME)

    print(f"md:   {md_desktop} ({md_desktop.stat().st_size} bytes)")
    print(f"docx: {docx_desktop} ({docx_desktop.stat().st_size} bytes)")


if __name__ == "__main__":
    main()
